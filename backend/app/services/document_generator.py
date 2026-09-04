import os
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FONT_DIR = os.path.join(BASE_DIR, "app", "fonts")
FONT_PATH = os.path.join(FONT_DIR, "Roboto-Regular.ttf")
FONT_BOLD_PATH = os.path.join(FONT_DIR, "Roboto-Bold.ttf")


class DocumentGenerator:

    # =========================================================================
    # СПРАВОЧНИКИ ДЛЯ КАРТЫ ОБРАБОТКИ ПДн
    # =========================================================================
    SYSTEM_TYPE_LABELS = {
        'local': 'Локальная система',
        'cloud_saas': 'Облачный SaaS',
        'file': 'Файл / таблица',
        'physical': 'Физический носитель',
    }

    CATEGORY_LABELS = {
        'employees': 'Сотрудники',
        'clients': 'Клиенты',
        'candidates': 'Кандидаты',
        'visitors': 'Посетители сайта',
        'contractors': 'Подрядчики',
    }

    # =========================================================================
    # СПРАВОЧНИК ТИПОВ ЗАПРОСОВ СУБЪЕКТОВ ПДн
    # =========================================================================
    REQUEST_TYPE_LABELS = {
        'information': 'о предоставлении информации об обработке персональных данных (ст. 14 Закона № 152-ФЗ)',
        'clarification': 'об уточнении (блокировании) персональных данных (ст. 20 Закона № 152-ФЗ)',
        'destruction': 'об уничтожении персональных данных (ст. 21 Закона № 152-ФЗ)',
        'withdrawal': 'о прекращении обработки персональных данных в связи с отзывом согласия',
    }

    # =========================================================================
    # ГЕНЕРАЦИЯ PDF
    # =========================================================================
    def _make_pdf(self, title: str, sections: list[dict]) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('Roboto', '', FONT_PATH, uni=True)
        pdf.add_font('Roboto', 'B', FONT_BOLD_PATH, uni=True)

        pdf.set_font('Roboto', 'B', 14)
        pdf.multi_cell(0, 8, title, align='C')
        pdf.ln(4)

        pdf.set_font('Roboto', '', 10)
        today = datetime.now().strftime('%d.%m.%Y')
        pdf.cell(0, 6, f"г. Москва, {today}", align='R', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        for section in sections:
            stype = section.get("type", "text")
            content = section.get("content", "")

            if stype == "subtitle":
                pdf.set_font('Roboto', 'B', 12)
                pdf.ln(2)
                pdf.multi_cell(0, 7, content)
                pdf.ln(1)
            elif stype == "text":
                pdf.set_font('Roboto', '', 11)
                pdf.multi_cell(0, 6, content)
                pdf.ln(1)
            elif stype == "item":
                pdf.set_font('Roboto', '', 11)
                pdf.set_x(15)
                pdf.multi_cell(0, 6, content)
                pdf.ln(0.5)
            elif stype == "sign":
                pdf.ln(4)
                pdf.set_font('Roboto', '', 11)
                pdf.multi_cell(0, 6, content)

        return pdf.output()

    # =========================================================================
    # ГЕНЕРАЦИЯ WORD
    # =========================================================================
    def _make_docx(self, title: str, sections: list[dict]) -> bytes:
        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

        today = datetime.now().strftime('%d.%m.%Y')
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"г. Москва, {today}")
        date_run.font.size = Pt(10)

        doc.add_paragraph()

        for section in sections:
            stype = section.get("type", "text")
            content = section.get("content", "")

            if stype == "subtitle":
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            elif stype == "text":
                p = doc.add_paragraph(content)
                p.paragraph_format.space_after = Pt(6)
            elif stype == "item":
                p = doc.add_paragraph(content)
                p.paragraph_format.left_indent = Cm(1.25)
                p.paragraph_format.space_after = Pt(3)
            elif stype == "sign":
                doc.add_paragraph()
                for line in content.split('\n'):
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(3)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =========================================================================
    # БЕЗОПАСНЫЕ ПРИМИТИВЫ (без multi_cell — без исключений fpdf2)
    # =========================================================================
    def _wrap_text(self, pdf, text, max_width):
        """Переносит текст по словам; слишком длинные слова режет по символам."""
        text = str(text) if text is not None else ''
        if text == '':
            return ['']

        lines = []
        for word in text.split(' '):
            # если слово шире колонки — режем по символам
            while pdf.get_string_width(word) > max_width and len(word) > 1:
                cut = len(word) - 1
                while cut > 1 and pdf.get_string_width(word[:cut]) > max_width:
                    cut -= 1
                lines.append(word[:cut])
                word = word[cut:]

            if not lines:
                lines.append(word)
            else:
                candidate = f"{lines[-1]} {word}"
                if pdf.get_string_width(candidate) <= max_width:
                    lines[-1] = candidate
                else:
                    lines.append(word)
        return lines

    def _write_wrapped(self, pdf, text, line_h, align='L'):
        """Вывод текста с переносом через cell() + автоперенос на новую страницу."""
        usable = pdf.w - pdf.l_margin - pdf.r_margin
        for raw_line in str(text).split('\n'):
            for line in self._wrap_text(pdf, raw_line, usable):
                if pdf.get_y() + line_h > 285:
                    pdf.add_page()
                pdf.cell(0, line_h, line, align=align, new_x="LMARGIN", new_y="NEXT")

    def _render_data_map_table(self, pdf, headers, rows, col_widths):
        """Таблица ИС: текст выводится через pdf.text() — без проверки ширины."""
        line_h = 4.5
        pad = 1.5
        x0 = 10  # левое поле FPDF по умолчанию
        total_w = sum(col_widths)

        def draw_header():
            pdf.set_font('Roboto', 'B', 8)
            pdf.set_fill_color(230, 230, 230)
            y = pdf.get_y()
            x = x0
            for i, h in enumerate(headers):
                pdf.rect(x, y, col_widths[i], 8, 'DF')
                pdf.set_xy(x, y + 2.5)
                pdf.cell(col_widths[i], 3, h, border=0, align='C')
                x += col_widths[i]
            pdf.set_xy(x0, y + 8)

        draw_header()

        pdf.set_font('Roboto', '', 8)
        for row in rows:
            # Перенос текста по колонкам
            wrapped = []
            max_lines = 1
            for i, text in enumerate(row):
                lines = self._wrap_text(pdf, text, col_widths[i] - 2 * pad)
                wrapped.append(lines)
                max_lines = max(max_lines, len(lines))
            row_h = max_lines * line_h + 2 * pad

            # Перенос на новую страницу, если строка не помещается
            if pdf.get_y() + row_h > 285:
                pdf.add_page()
                draw_header()
                pdf.set_font('Roboto', '', 8)

            y = pdf.get_y()

            # Рамка строки и вертикальные разделители
            pdf.rect(x0, y, total_w, row_h)
            x = x0
            for w in col_widths[:-1]:
                x += w
                pdf.line(x, y, x, y + row_h)

            # Текст ячеек
            x = x0
            for i, lines in enumerate(wrapped):
                for li, ln in enumerate(lines):
                    pdf.text(x + pad, y + pad + line_h * (li + 0.8), ln)
                x += col_widths[i]

            pdf.set_xy(x0, y + row_h)

    def generate_data_map_pdf(self, company_data: dict, data_systems: list) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')

        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('Roboto', '', FONT_PATH, uni=True)
        pdf.add_font('Roboto', 'B', FONT_BOLD_PATH, uni=True)

        # Шапка документа
        pdf.set_font('Roboto', 'B', 14)
        self._write_wrapped(pdf, "КАРТА ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ", 8, align='C')
        self._write_wrapped(pdf, name, 8, align='C')
        pdf.ln(2)

        pdf.set_font('Roboto', '', 10)
        today = datetime.now().strftime('%d.%m.%Y')
        pdf.cell(0, 6, f"Дата формирования: {today}", align='R', new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        self._write_wrapped(pdf, f"Оператор: {name} (ИНН: {inn})", 6)
        self._write_wrapped(
            pdf,
            "Настоящий документ определяет информационные системы (ИС), в которых оператором "
            "обрабатываются персональные данные, и составлен в соответствии с требованиями "
            "Федерального закона от 27.07.2006 № 152-ФЗ «О персональных данных».",
            6,
        )
        pdf.ln(4)

        if not data_systems:
            pdf.set_font('Roboto', '', 11)
            self._write_wrapped(
                pdf,
                "На дату формирования документа информационные системы, обрабатывающие персональные данные, у оператора не зарегистрированы.",
                6,
            )
        else:
            headers = ["№", "Информационная система", "Категории субъектов", "Локация данных", "Ответственный", "Субъектов"]
            col_widths = [8, 50, 38, 42, 40, 12]

            rows = []
            total_subjects = 0
            for idx, ds in enumerate(data_systems, start=1):
                type_label = self.SYSTEM_TYPE_LABELS.get(ds.get('system_type', ''), ds.get('system_type', ''))
                cats = ", ".join(
                    self.CATEGORY_LABELS.get(c, c) for c in ds.get('categories', [])
                ) or "—"
                location = ds.get('data_location') or "—"
                resp_parts = []
                if ds.get('responsible_name'):
                    resp_parts.append(ds['responsible_name'])
                if ds.get('responsible_position'):
                    resp_parts.append(f"({ds['responsible_position']})")
                resp = " ".join(resp_parts) or "—"
                subj = ds.get('pd_subjects_count', 0) or 0
                total_subjects += subj

                rows.append([
                    str(idx),
                    f"{ds.get('name', '')} — {type_label}",
                    cats,
                    location,
                    resp,
                    str(subj),
                ])

            self._render_data_map_table(pdf, headers, rows, col_widths)

            pdf.ln(4)
            pdf.set_font('Roboto', 'B', 10)
            self._write_wrapped(
                pdf,
                f"Итого информационных систем: {len(data_systems)}. "
                f"Всего субъектов ПДн в системах: {total_subjects}.",
                6,
            )

        # Блок подписи
        pdf.ln(6)
        pdf.set_font('Roboto', '', 11)
        self._write_wrapped(pdf, "Ответственный за организацию обработки персональных данных:", 6)
        pdf.ln(4)
        pdf.cell(0, 6, "_________________ / _________________________________", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, "(подпись)                          (расшифровка подписи)", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.cell(0, 6, "Дата: «___» _______________ 20___ г.", new_x="LMARGIN", new_y="NEXT")

        return pdf.output()

    # =========================================================================
    # ОТВЕТ НА ЗАПРОС СУБЪЕКТА ПДн (киллер-фича)
    # =========================================================================
    def generate_subject_response_pdf(self, company_data: dict, request_data: dict, subject_data: dict = None) -> bytes:
        """Генерирует юридически корректный ответ на запрос субъекта ПДн.

        company_data: name, inn, address, director_name
        request_data: subject_name, request_type, received_at, deadline (строки дат dd.mm.yyyy)
        subject_data: None или dict: category, legal_basis, data_types, data_systems (list)
        """
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        address = company_data.get('address') or ''
        subject_name = request_data.get('subject_name', '')
        rtype = request_data.get('request_type', 'information')
        received = request_data.get('received_at', '')
        deadline = request_data.get('deadline', '')
        type_label = self.REQUEST_TYPE_LABELS.get(rtype, self.REQUEST_TYPE_LABELS['information'])

        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('Roboto', '', FONT_PATH, uni=True)
        pdf.add_font('Roboto', 'B', FONT_BOLD_PATH, uni=True)

        # Шапка: реквизиты оператора
        pdf.set_font('Roboto', 'B', 13)
        self._write_wrapped(pdf, name, 7, align='C')
        pdf.set_font('Roboto', '', 9)
        self._write_wrapped(pdf, f"ИНН: {inn}" + (f" · {address}" if address else ""), 5, align='C')
        pdf.ln(4)

        # Заголовок
        pdf.set_font('Roboto', 'B', 12)
        self._write_wrapped(pdf, "ОТВЕТ НА ЗАПРОС СУБЪЕКТА ПЕРСОНАЛЬНЫХ ДАННЫХ", 7, align='C')
        pdf.set_font('Roboto', '', 10)
        self._write_wrapped(pdf, f"({type_label})", 6, align='C')
        pdf.ln(3)

        # Дата и адресат
        today = datetime.now().strftime('%d.%m.%Y')
        pdf.set_font('Roboto', '', 10)
        pdf.cell(0, 6, f"Дата ответа: {today}", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, f"Кому: {subject_name}", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, f"Запрос получен: {received} · Срок ответа: {deadline}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # Преамбула
        self._write_wrapped(
            pdf,
            f"Уважаемый(ая) {subject_name}! В ответ на Ваш запрос, полученный Оператором {received}, "
            "настоящим сообщаем следующее в соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ «О персональных данных».",
            6,
        )
        pdf.ln(2)

        # Тело ответа по типу запроса
        pdf.set_font('Roboto', '', 10)
        if rtype == 'information':
            legal = (subject_data or {}).get('legal_basis')
            data_types = (subject_data or {}).get('data_types')
            body = [
                "1. Оператор подтверждает факт обработки Ваших персональных данных.",
                "2. Цели обработки: исполнение договора, соблюдение требований законодательства РФ, а также иные цели, указанные в согласии на обработку персональных данных.",
                f"3. Правовые основания обработки: {legal or 'согласие субъекта, договор, требования законодательства РФ'}.",
                f"4. Категории обрабатываемых персональных данных: {data_types or 'ФИО и сведения, необходимые для исполнения договора и требований законодательства'}.",
                "5. Обработка осуществляется как с использованием средств автоматизации, так и без них.",
                "6. Ваши персональные данные не передаются третьим лицам, за исключением случаев, предусмотренных законодательством РФ.",
            ]
        elif rtype == 'clarification':
            body = [
                "1. Ваш запрос об уточнении (блокировании) персональных данных принят Оператором в работу.",
                "2. В соответствии со ст. 20 Закона № 152-ФЗ Оператор в течение 7 рабочих дней со дня представления подтверждающих документов осуществит блокирование и (или) уточнение персональных данных.",
                "3. Просим приложить документы, подтверждающие необходимость уточнения данных.",
                "4. О результатах рассмотрения Вы будете уведомлены в установленный законом срок.",
            ]
        elif rtype == 'destruction':
            body = [
                "1. Ваш запрос об уничтожении персональных данных принят Оператором в работу.",
                "2. В соответствии со ст. 21 Закона № 152-ФЗ при подтверждении оснований, предусмотренных законом, Оператор прекратит обработку и уничтожит Ваши персональные данные в срок, не превышающий 30 дней с даты поступления запроса.",
                "3. Обращаем внимание: уничтожение данных не осуществляется, если Оператор обязан продолжать обработку на основании закона (требования трудового, налогового законодательства о хранении документов).",
                "4. О результатах рассмотрения Вы будете уведомлены в письменной форме.",
            ]
        else:  # withdrawal
            body = [
                "1. Ваш отзыв согласия на обработку персональных данных принят Оператором.",
                "2. В соответствии с ч. 5 ст. 9 Закона № 152-ФЗ Оператор прекратит обработку и уничтожит Ваши персональные данные в срок, не превышающий 30 дней с даты поступления отзыва, если иное не предусмотрено договором или законом.",
                "3. Оператор вправе продолжить обработку без согласия при наличии оснований, указанных в п. 2–11 ч. 1 ст. 6, ч. 2 ст. 10 и ч. 2 ст. 11 Закона № 152-ФЗ.",
                "4. О результатах рассмотрения Вы будете уведомлены в письменной форме.",
            ]

        for line in body:
            self._write_wrapped(pdf, line, 6)
            pdf.ln(1)

        # Блок информационных систем (если есть связь с Реестром)
        systems = (subject_data or {}).get('data_systems') or []
        pdf.ln(1)
        if systems:
            pdf.set_font('Roboto', 'B', 10)
            self._write_wrapped(pdf, "Информационные системы, в которых обрабатываются Ваши персональные данные:", 6)
            pdf.set_font('Roboto', '', 10)
            for s in systems:
                t = self.SYSTEM_TYPE_LABELS.get(s.get('system_type', ''), s.get('system_type', ''))
                loc = s.get('data_location')
                line = f"— {s.get('name', '')} ({t})" + (f", место хранения: {loc}" if loc else "")
                self._write_wrapped(pdf, line, 6)
        else:
            self._write_wrapped(
                pdf,
                "Перечень информационных систем, в которых обрабатываются Ваши персональные данные, приведён в Карте обработки персональных данных Оператора.",
                6,
            )

        # Подпись
        pdf.ln(6)
        pdf.set_font('Roboto', '', 10)
        self._write_wrapped(pdf, "Ответственный за организацию обработки персональных данных:", 6)
        pdf.ln(3)
        pdf.cell(0, 6, "_________________ / _________________________________", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, "(подпись)                          (расшифровка подписи)", new_x="LMARGIN", new_y="NEXT")

        return pdf.output()

    # =========================================================================
    # ОТЧЁТ О ПРОВЕРКЕ САЙТА НА СООТВЕТСТВИЕ 152-ФЗ
    # =========================================================================
    def generate_compliance_report_pdf(self, company_data: dict, check_data: dict) -> bytes:
        """PDF-отчёт по результатам автоматической проверки сайта.

        company_data: name, inn (может быть пустым — отчёт по «голому» URL)
        check_data: url, checked_at, compliance_percentage, total_required,
                    passed_required, checks {key: {name, required, found, details[]}}
        """
        name = (company_data or {}).get('name') or ''
        inn = (company_data or {}).get('inn') or ''

        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('Roboto', '', FONT_PATH, uni=True)
        pdf.add_font('Roboto', 'B', FONT_BOLD_PATH, uni=True)

        # Шапка
        pdf.set_font('Roboto', 'B', 14)
        self._write_wrapped(pdf, "ОТЧЁТ О ПРОВЕРКЕ САЙТА", 8, align='C')
        self._write_wrapped(pdf, "НА СООТВЕТСТВИЕ ТРЕБОВАНИЯМ 152-ФЗ", 8, align='C')
        pdf.ln(3)

        pdf.set_font('Roboto', '', 10)
        checked_at = check_data.get('checked_at', '')
        try:
            checked_str = datetime.fromisoformat(str(checked_at).replace('Z', '+00:00')).strftime('%d.%m.%Y %H:%M')
        except Exception:
            checked_str = str(checked_at)

        self._write_wrapped(pdf, f"Проверяемый сайт: {check_data.get('url', '')}", 6)
        self._write_wrapped(pdf, f"Дата проверки: {checked_str}", 6)
        if name:
            self._write_wrapped(pdf, f"Оператор: {name}" + (f" (ИНН: {inn})" if inn else ""), 6)
        pdf.ln(3)

        # Крупный балл соответствия
        score = int(check_data.get('compliance_percentage', 0) or 0)
        passed = check_data.get('passed_required', 0)
        total = check_data.get('total_required', 0)

        if score >= 80:
            pdf.set_text_color(0, 150, 60)
        elif score >= 60:
            pdf.set_text_color(200, 150, 0)
        else:
            pdf.set_text_color(220, 50, 50)
        pdf.set_font('Roboto', 'B', 26)
        pdf.cell(0, 12, f"{score}%", align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Roboto', '', 10)
        self._write_wrapped(pdf, f"Соответствие обязательным требованиям ({passed} из {total} пунктов)", 6, align='C')
        pdf.ln(4)

        # Детальные результаты
        checks = check_data.get('checks', {}) or {}
        pdf.set_font('Roboto', 'B', 11)
        self._write_wrapped(pdf, "РЕЗУЛЬТАТЫ ПРОВЕРКИ:", 7)
        pdf.ln(1)

        for _key, ch in checks.items():
            found = bool(ch.get('found'))
            required = bool(ch.get('required'))

            if found:
                mark, color = "ВЫПОЛНЕНО", (0, 150, 60)
            elif required:
                mark, color = "НЕ ВЫПОЛНЕНО", (220, 50, 50)
            else:
                mark, color = "РЕКОМЕНДАЦИЯ", (200, 150, 0)

            pdf.set_text_color(*color)
            pdf.set_font('Roboto', 'B', 10)
            self._write_wrapped(pdf, f"[{mark}] {ch.get('name', '')}", 6)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('Roboto', '', 9)
            self._write_wrapped(pdf, f"Приоритет: {'Обязательно' if required else 'Рекомендуется'}", 5)
            for d in (ch.get('details') or [])[:5]:
                self._write_wrapped(pdf, f"• {d}", 5)
            pdf.ln(2)

        # Критические пробелы
        missing = [ch for ch in checks.values() if ch.get('required') and not ch.get('found')]
        pdf.set_font('Roboto', 'B', 11)
        if missing:
            pdf.set_text_color(220, 50, 50)
            self._write_wrapped(pdf, f"КРИТИЧЕСКИЕ ПРОБЕЛЫ ({len(missing)}):", 7)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('Roboto', '', 10)
            for ch in missing:
                self._write_wrapped(pdf, f"• {ch.get('name', '')}", 6)
            pdf.ln(2)
            self._write_wrapped(
                pdf,
                "Рекомендация: устранить критические пробелы до проверки Роскомнадзора. "
                "Штраф за отсутствие Политики обработки ПДн — до 60 000 ₽, за отсутствие согласия — до 300 000 ₽.",
                6,
            )
        else:
            pdf.set_text_color(0, 150, 60)
            self._write_wrapped(pdf, "Критических пробелов не выявлено: все обязательные требования выполнены.", 7)
            pdf.set_text_color(0, 0, 0)

        # Дисклеймер
        pdf.ln(4)
        pdf.set_font('Roboto', '', 8)
        self._write_wrapped(
            pdf,
            "Юридический дисклеймер: результаты автоматической проверки носят исключительно информационный "
            "характер и не являются официальным юридическим заключением. Для получения юридически значимого "
            "аудита соответствия 152-ФЗ рекомендуем обратиться к профильным специалистам.",
            5,
        )

        # Подпись
        pdf.ln(4)
        pdf.set_font('Roboto', '', 10)
        pdf.cell(0, 6, "_________________ / _________________________________", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, "(подпись)                          (расшифровка подписи)", new_x="LMARGIN", new_y="NEXT")

        return pdf.output()

    # =========================================================================
    # ДАННЫЕ ДЛЯ ДОКУМЕНТОВ (ОБЩИЕ)
    # =========================================================================
    def _policy_sections(self, name, inn, email):
        return [
            {"type": "subtitle", "content": "1. ОБЩИЕ ПОЛОЖЕНИЯ"},
            {"type": "item", "content": f"1.1. Настоящая Политика разработана в соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ «О персональных данных» (далее — Закон № 152-ФЗ) и определяет порядок обработки персональных данных и меры по обеспечению их безопасности, предпринимаемые {name} (далее — Оператор)."},
            {"type": "item", "content": "1.2. Оператор ставит своей важнейшей целью и условием осуществления своей деятельности соблюдение прав и свобод человека и гражданина при обработке его персональных данных, в том числе защиты прав на неприкосновенность частной жизни, личную и семейную тайну."},
            {"type": "item", "content": "1.3. Настоящая Политика применяется ко всей информации, которую Оператор может получить о посетителях веб-сайта, клиентах, сотрудниках и контрагентах."},
            {"type": "subtitle", "content": "2. ЦЕЛИ ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ"},
            {"type": "item", "content": "2.1. Оператор обрабатывает персональные данные в следующих целях:"},
            {"type": "item", "content": "— исполнение договоров, стороной которых является субъект персональных данных;"},
            {"type": "item", "content": "— предоставление услуг и сервисов клиентам Оператора;"},
            {"type": "item", "content": "— трудоустройство и кадровое делопроизводство;"},
            {"type": "item", "content": "— направление информационных и рекламных сообщений (с согласия субъекта);"},
            {"type": "item", "content": "— соблюдение требований законодательства Российской Федерации."},
            {"type": "subtitle", "content": "3. ПРАВОВЫЕ ОСНОВАНИЯ ОБРАБОТКИ"},
            {"type": "item", "content": "3.1. Правовыми основаниями обработки персональных данных Оператором являются:"},
            {"type": "item", "content": "— Конституция Российской Федерации;"},
            {"type": "item", "content": "— Федеральный закон № 152-ФЗ «О персональных данных»;"},
            {"type": "item", "content": "— Трудовой кодекс Российской Федерации;"},
            {"type": "item", "content": "— согласие субъекта персональных данных на обработку его персональных данных."},
            {"type": "subtitle", "content": "4. ОБЪЕМ И КАТЕГОРИИ ОБРАБАТЫВАЕМЫХ ДАННЫХ"},
            {"type": "item", "content": "4.1. Оператор может обрабатывать следующие персональные данные:"},
            {"type": "item", "content": "— фамилия, имя, отчество;"},
            {"type": "item", "content": "— дата и место рождения;"},
            {"type": "item", "content": "— адрес электронной почты, номер телефона;"},
            {"type": "item", "content": "— паспортные данные;"},
            {"type": "item", "content": "— сведения о трудовой деятельности."},
            {"type": "item", "content": "4.2. Оператор не обрабатывает специальные категории персональных данных, за исключением случаев, прямо предусмотренных законом."},
            {"type": "subtitle", "content": "5. ПОРЯДОК И УСЛОВИЯ ОБРАБОТКИ"},
            {"type": "item", "content": "5.1. Обработка персональных данных осуществляется с согласия субъекта, за исключением случаев, предусмотренных Законом № 152-ФЗ."},
            {"type": "item", "content": "5.2. Оператор не раскрывает третьим лицам и не распространяет персональные данные без согласия субъекта, если иное не предусмотрено федеральным законом."},
            {"type": "item", "content": "5.3. Оператор принимает необходимые правовые, организационные и технические меры для защиты персональных данных."},
            {"type": "subtitle", "content": "6. ПРАВА СУБЪЕКТА ПЕРСОНАЛЬНЫХ ДАННЫХ"},
            {"type": "item", "content": "6.1. Субъект персональных данных имеет право:"},
            {"type": "item", "content": "— получать сведения об обработке своих персональных данных;"},
            {"type": "item", "content": "— требовать уточнения, блокирования или уничтожения своих данных;"},
            {"type": "item", "content": "— отозвать согласие на обработку персональных данных;"},
            {"type": "item", "content": "— обжаловать действия Оператора в Роскомнадзор или в суд."},
            {"type": "subtitle", "content": "7. МЕРЫ ПО ОБЕСПЕЧЕНИЮ БЕЗОПАСНОСТИ"},
            {"type": "item", "content": "7.1. Оператор применяет следующие меры защиты:"},
            {"type": "item", "content": "— назначение ответственного за организацию обработки персональных данных;"},
            {"type": "item", "content": "— ограничение доступа к персональным данным;"},
            {"type": "item", "content": "— использование средств защиты информации;"},
            {"type": "item", "content": "— регулярный контроль эффективности принимаемых мер."},
            {"type": "subtitle", "content": "8. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ"},
            {"type": "item", "content": "8.1. Настоящая Политика действует бессрочно до момента её отзыва Оператором."},
            {"type": "item", "content": "8.2. Оператор вправе вносить изменения в настоящую Политику. Новая редакция вступает в силу с момента её размещения на официальном сайте Оператора."},
            {"type": "sign", "content": f"Оператор: {name}\nИНН: {inn}\nEmail: {email}\n\nГенеральный директор _________________ / _________________"},
        ]

    def _consent_sections(self, name, inn, email):
        return [
            {"type": "text", "content": f"Я, _________________________________________________________________________________,"},
            {"type": "text", "content": "(фамилия, имя, отчество полностью)"},
            {"type": "text", "content": "паспорт серии _______ № ____________, выдан ___________________________________________________"},
            {"type": "text", "content": "_________________________________________________________________________________________________,"},
            {"type": "text", "content": "(кем и когда выдан)"},
            {"type": "text", "content": f"зарегистрированный(ая) по адресу: _______________________________________________________________"},
            {"type": "text", "content": f"в целях заключения и исполнения договоров с оператором {name},"},
            {"type": "text", "content": "даю согласие на обработку моих персональных данных со следующими условиями:"},
            {"type": "subtitle", "content": "1. Перечень персональных данных, на обработку которых дается согласие:"},
            {"type": "item", "content": "— фамилия, имя, отчество;"},
            {"type": "item", "content": "— дата и место рождения;"},
            {"type": "item", "content": "— пол;"},
            {"type": "item", "content": "— паспортные данные (серия, номер, кем и когда выдан);"},
            {"type": "item", "content": "— адрес регистрации и фактического проживания;"},
            {"type": "item", "content": "— контактный телефон, адрес электронной почты;"},
            {"type": "item", "content": "— сведения об образовании, профессии, месте работы;"},
            {"type": "item", "content": "— ИНН, СНИЛС;"},
            {"type": "item", "content": "— иные данные, необходимые для исполнения договора."},
            {"type": "subtitle", "content": "2. Цели обработки персональных данных:"},
            {"type": "item", "content": "— заключение и исполнение договоров с оператором;"},
            {"type": "item", "content": "— предоставление услуг и сервисов;"},
            {"type": "item", "content": "— направление информационных и рекламных сообщений;"},
            {"type": "item", "content": "— исполнение требований законодательства Российской Федерации."},
            {"type": "subtitle", "content": "3. Перечень действий с персональными данными:"},
            {"type": "item", "content": "Согласие дается на обработку персональных данных как с использованием средств автоматизации, так и без их использования, включая:"},
            {"type": "item", "content": "— сбор, запись, систематизацию, накопление, хранение;"},
            {"type": "item", "content": "— уточнение (обновление, изменение), извлечение;"},
            {"type": "item", "content": "— использование, передачу (предоставление, доступ);"},
            {"type": "item", "content": "— обезличивание, блокирование, удаление, уничтожение."},
            {"type": "subtitle", "content": "4. Срок обработки и порядок отзыва согласия:"},
            {"type": "item", "content": "4.1. Настоящее согласие действует с момента его подписания до момента отзыва субъектом персональных данных."},
            {"type": "item", "content": f"4.2. Согласие может быть отозвано субъектом персональных данных путем направления письменного уведомления оператору по адресу: {email}."},
            {"type": "item", "content": "4.3. В случае отзыва согласия оператор вправе продолжить обработку персональных данных без согласия субъекта при наличии оснований, указанных в п. 2-11 ч. 1 ст. 6, ч. 2 ст. 10 и ч. 2 ст. 11 Федерального закона № 152-ФЗ."},
            {"type": "subtitle", "content": "5. Подтверждение:"},
            {"type": "item", "content": "Настоящим я подтверждаю, что:"},
            {"type": "item", "content": "— мне разъяснены цели, способы и условия обработки моих персональных данных;"},
            {"type": "item", "content": "— я ознакомлен(а) с Политикой в отношении обработки персональных данных оператора;"},
            {"type": "item", "content": "— я понимаю последствия предоставления и отзыва настоящего согласия;"},
            {"type": "item", "content": "— я даю согласие добровольно, в полном объеме осознавая значение своих действий."},
            {"type": "sign", "content": f"Оператор: {name} (ИНН {inn})\n\nДата: «___» _______________ 20___ г.\n\nПодпись субъекта: _________________ / _________________________________\n(подпись)                          (расшифровка подписи)"},
        ]

    def _nda_sections(self, name, inn, email):
        return [
            {"type": "text", "content": f"Я, нижеподписавшийся(аяся) _________________________________________________________________,"},
            {"type": "text", "content": "(фамилия, имя, отчество полностью)"},
            {"type": "text", "content": f"работающий(ая) в {name} (далее — Работодатель) в должности _______________________,"},
            {"type": "text", "content": "в соответствии со статьей 139 Трудового кодекса Российской Федерации и Федеральным законом от 27.07.2006 № 152-ФЗ «О персональных данных»,"},
            {"type": "text", "content": "принимаю на себя следующие обязательства:"},
            {"type": "subtitle", "content": "1. Предмет обязательства"},
            {"type": "item", "content": "1.1. Я обязуюсь не разглашать ставшие мне известными в процессе трудовой деятельности персональные данные работников, клиентов, контрагентов и иных субъектов персональных данных, обрабатываемых Работодателем."},
            {"type": "item", "content": "1.2. Под разглашением понимается действие или бездействие, в результате которых персональные данные становятся известными третьим лицам без согласия субъекта персональных данных."},
            {"type": "subtitle", "content": "2. Обязанности работника"},
            {"type": "item", "content": "2.1. Не сообщать персональные данные третьим лицам, а также не раскрывать их без письменного согласия субъекта персональных данных или Работодателя."},
            {"type": "item", "content": "2.2. Не использовать персональные данные в целях, не связанных с исполнением трудовых обязанностей."},
            {"type": "item", "content": "2.3. Принимать меры по защите персональных данных от несанкционированного доступа."},
            {"type": "item", "content": "2.4. Немедленно сообщать Работодателю о любых случаях утечки или возможной утечки персональных данных."},
            {"type": "item", "content": "2.5. Соблюдать требования внутренних документов Работодателя в области обработки персональных данных."},
            {"type": "subtitle", "content": "3. Ответственность"},
            {"type": "item", "content": "3.1. За разглашение персональных данных я несу дисциплинарную, гражданско-правовую, административную и уголовную ответственность в соответствии с действующим законодательством Российской Федерации."},
            {"type": "item", "content": "3.2. Работодатель вправе применить ко мне меры дисциплинарного взыскания вплоть до увольнения по пп. «в» п. 6 ч. 1 ст. 81 Трудового кодекса Российской Федерации."},
            {"type": "subtitle", "content": "4. Срок действия обязательства"},
            {"type": "item", "content": "4.1. Настоящее обязательство действует в течение всего срока трудовой деятельности у Работодателя."},
            {"type": "item", "content": "4.2. Обязательство о неразглашении сохраняется в силе в течение 3 (трёх) лет после прекращения трудового договора."},
            {"type": "subtitle", "content": "5. Заключительные положения"},
            {"type": "item", "content": "5.1. Настоящее обязательство составлено в двух экземплярах, имеющих одинаковую юридическую силу: один — у Работодателя, второй — у работника."},
            {"type": "item", "content": "5.2. Обязательство вступает в силу с момента его подписания работником."},
            {"type": "sign", "content": f"Работодатель: {name}\nИНН: {inn}\nEmail: {email}\n\nДата: «___» _______________ 20___ г.\n\nПодпись работника: _________________ / _________________________________\n(подпись)                          (расшифровка подписи)"},
        ]

    def _order_sections(self, name, inn, email):
        return [
            {"type": "text", "content": f"В целях обеспечения соблюдения требований Федерального закона от 27.07.2006 № 152-ФЗ «О персональных данных», а также для организации обработки персональных данных и обеспечения их безопасности,"},
            {"type": "text", "content": "ПРИКАЗЫВАЮ:"},
            {"type": "subtitle", "content": "1. Назначить ответственным"},
            {"type": "item", "content": f"1.1. Назначить ответственным за организацию обработки персональных данных в {name} _________________________________________________________________"},
            {"type": "item", "content": "(должность, фамилия, имя, отчество)"},
            {"type": "item", "content": "с «___» _______________ 20___ г."},
            {"type": "subtitle", "content": "2. Возложить обязанности"},
            {"type": "item", "content": "2.1. На ответственного возлагаются следующие обязанности:"},
            {"type": "item", "content": "— принятие и рассмотрение обращений и запросов субъектов персональных данных или их представителей;"},
            {"type": "item", "content": "— организация обработки персональных данных в соответствии с требованиями Закона № 152-ФЗ;"},
            {"type": "item", "content": "— контроль за соблюдением требований законодательства в области персональных данных;"},
            {"type": "item", "content": "— доведение до сведения работников организации требований законодательства и внутренних документов по обработке персональных данных;"},
            {"type": "item", "content": "— организация обучения работников в области обработки персональных данных;"},
            {"type": "item", "content": "— взаимодействие с уполномоченным органом по защите прав субъектов персональных данных (Роскомнадзор);"},
            {"type": "item", "content": "— внутренний контроль за соблюдением условий обработки персональных данных."},
            {"type": "subtitle", "content": "3. Обеспечить доступ к информации"},
            {"type": "item", "content": "3.1. Обеспечить неограниченный доступ к настоящему приказу и Политике в отношении обработки персональных данных."},
            {"type": "subtitle", "content": "4. Контроль"},
            {"type": "item", "content": "4.1. Контроль за исполнением настоящего приказа оставляю за собой."},
            {"type": "sign", "content": f"Основание: Федеральный закон от 27.07.2006 № 152-ФЗ «О персональных данных».\n\nГенеральный директор {name}\n\n_______________ / _________________________________\n(подпись)      (расшифровка подписи)\n\nС приказом ознакомлен(а):\n\n_______________ / _________________________________\n(подпись)      (расшифровка подписи)\n\n«___» _______________ 20___ г."},
        ]

    def _threat_sections(self, name, inn, email):
        return [
            {"type": "subtitle", "content": "1. ОБЩИЕ СВЕДЕНИЯ"},
            {"type": "item", "content": f"1.1. Настоящая Модель угроз разработана в соответствии с требованиями Федерального закона от 27.07.2006 № 152-ФЗ «О персональных данных» и методическими документами ФСТЭК России."},
            {"type": "item", "content": f"1.2. Оператор персональных данных: {name}, ИНН {inn}."},
            {"type": "item", "content": f"1.3. Контактное лицо: {email}."},
            {"type": "item", "content": "1.4. Объектом защиты являются персональные данные, обрабатываемые в информационных системах оператора."},
            {"type": "subtitle", "content": "2. ХАРАКТЕРИСТИКА ОБЪЕКТА ЗАЩИТЫ"},
            {"type": "item", "content": "2.1. Категории обрабатываемых персональных данных:"},
            {"type": "item", "content": "— общие персональные данные (ФИО, адрес, телефон, email);"},
            {"type": "item", "content": "— специальные персональные данные (при наличии согласия субъекта);"},
            {"type": "item", "content": "— биометрические персональные данные (не обрабатываются)."},
            {"type": "item", "content": "2.2. Количество субъектов персональных данных: до 1000 человек."},
            {"type": "item", "content": "2.3. Уровень защищённости информационной системы персональных данных (ИСПДн): устанавливается по результатам классификации."},
            {"type": "subtitle", "content": "3. АНАЛИЗ УГРОЗ БЕЗОПАСНОСТИ"},
            {"type": "item", "content": "3.1. Актуальные угрозы безопасности персональных данных:"},
            {"type": "item", "content": "— угрозы 1-го типа (внешний нарушитель с низким уровнем возможностей):"},
            {"type": "item", "content": "  • несанкционированный доступ через интернет;"},
            {"type": "item", "content": "  • фишинговые атаки;"},
            {"type": "item", "content": "  • вредоносное программное обеспечение."},
            {"type": "item", "content": "— угрозы 2-го типа (внутренний нарушитель с привилегированным доступом):"},
            {"type": "item", "content": "  • умышленное копирование персональных данных;"},
            {"type": "item", "content": "  • модификация или уничтожение данных;"},
            {"type": "item", "content": "  • передача данных третьим лицам."},
            {"type": "item", "content": "— угрозы 3-го типа (внешний нарушитель с высоким уровнем возможностей):"},
            {"type": "item", "content": "  • целевые атаки на информационные системы;"},
            {"type": "item", "content": "  • перехват данных при передаче по сетям."},
            {"type": "subtitle", "content": "4. МЕРЫ ПО ЗАЩИТЕ ПЕРСОНАЛЬНЫХ ДАННЫХ"},
            {"type": "item", "content": "4.1. Организационные меры:"},
            {"type": "item", "content": "— назначение ответственного за обработку персональных данных;"},
            {"type": "item", "content": "— разработка и утверждение Политики обработки персональных данных;"},
            {"type": "item", "content": "— подписание обязательств о неразглашении с работниками;"},
            {"type": "item", "content": "— регулярное обучение работников."},
            {"type": "item", "content": "4.2. Технические меры:"},
            {"type": "item", "content": "— разграничение прав доступа к персональным данным;"},
            {"type": "item", "content": "— использование средств криптографической защиты информации;"},
            {"type": "item", "content": "— антивирусная защита;"},
            {"type": "item", "content": "— межсетевое экранирование;"},
            {"type": "item", "content": "— резервное копирование данных;"},
            {"type": "item", "content": "— журналирование событий безопасности."},
            {"type": "subtitle", "content": "5. ОЦЕНКА ЭФФЕКТИВНОСТИ МЕР ЗАЩИТЫ"},
            {"type": "item", "content": "5.1. Оценка эффективности мер защиты проводится не реже одного раза в три года, а также при изменении условий обработки персональных данных."},
            {"type": "item", "content": "5.2. По результатам оценки вносятся необходимые изменения в Модель угроз и систему защиты."},
            {"type": "subtitle", "content": "6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ"},
            {"type": "item", "content": "6.1. Настоящая Модель угроз подлежит пересмотру при изменении законодательства, структуры информационных систем или появлении новых актуальных угроз."},
            {"type": "sign", "content": f"Оператор: {name}\nИНН: {inn}\nEmail: {email}\n\nОтветственный за обработку ПДн: _________________ / _________________________________\n(подпись)                                        (расшифровка подписи)\n\nДата: «___» _______________ 20___ г."},
        ]

    # =========================================================================
    # ПУБЛИЧНЫЕ МЕТОДЫ: PDF
    # =========================================================================
    def generate_policy_152fz(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"ПОЛИТИКА В ОТНОШЕНИИ ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ\n{name}"
        return self._make_pdf(title, self._policy_sections(name, inn, email))

    def generate_consent_152fz(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"СОГЛАСИЕ СУБЪЕКТА НА ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ\nв соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ"
        return self._make_pdf(title, self._consent_sections(name, inn, email))

    def generate_nda_152fz(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"ОБЯЗАТЕЛЬСТВО О НЕРАЗГЛАШЕНИИ ПЕРСОНАЛЬНЫХ ДАННЫХ\nсотрудника {name}"
        return self._make_pdf(title, self._nda_sections(name, inn, email))

    def generate_order_responsible_152fz(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"ПРИКАЗ\nо назначении ответственного за организацию обработки персональных данных\n{name}"
        return self._make_pdf(title, self._order_sections(name, inn, email))

    def generate_threat_model_fstek(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"МОДЕЛЬ УГРОЗ БЕЗОПАСНОСТИ ПЕРСОНАЛЬНЫХ ДАННЫХ\n{name}"
        return self._make_pdf(title, self._threat_sections(name, inn, email))

    # =========================================================================
    # ПУБЛИЧНЫЕ МЕТОДЫ: WORD
    # =========================================================================
    def generate_policy_152fz_word(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"ПОЛИТИКА В ОТНОШЕНИИ ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ\n{name}"
        return self._make_docx(title, self._policy_sections(name, inn, email))

    def generate_consent_152fz_word(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"СОГЛАСИЕ СУБЪЕКТА НА ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ\nв соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ"
        return self._make_docx(title, self._consent_sections(name, inn, email))

    def generate_nda_152fz_word(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"ОБЯЗАТЕЛЬСТВО О НЕРАЗГЛАШЕНИИ ПЕРСОНАЛЬНЫХ ДАННЫХ\nсотрудника {name}"
        return self._make_docx(title, self._nda_sections(name, inn, email))

    def generate_order_responsible_152fz_word(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"ПРИКАЗ\nо назначении ответственного за организацию обработки персональных данных\n{name}"
        return self._make_docx(title, self._order_sections(name, inn, email))

    def generate_threat_model_fstek_word(self, company_data: dict) -> bytes:
        name = company_data.get('name', 'Организация')
        inn = company_data.get('inn', '')
        email = company_data.get('email', '')
        title = f"МОДЕЛЬ УГРОЗ БЕЗОПАСНОСТИ ПЕРСОНАЛЬНЫХ ДАННЫХ\n{name}"
        return self._make_docx(title, self._threat_sections(name, inn, email))


document_generator = DocumentGenerator()